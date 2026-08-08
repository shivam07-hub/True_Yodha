"""
cv_parser.py
Real CV → Lightcast skill extractor.

Pipeline:
  1. Extract text from PDF (pymupdf) or DOCX (python-docx)
  2. LLM extraction via OpenRouter — returns Lightcast-shaped skill names
  3. Validate each taxonomy_key against lightcast_skills_taxonomy.json
     (exact match → fuzzy fallback via difflib if no exact hit)
  4. Return signals matching the contract consumed by the scoring orchestrator:
        {
          "skills_detected": [
            {"taxonomy_key", "xp_awarded", "signal_type", "evidence"}
          ],
          "raw_text": str
        }

Signal types and XP (unchanged from v1 stub contract):
  "mention"    → +50   (named in skills section only)
  "project"    → +150  (used in a real project or role)
  "impact"     → +350  (applied with measurable scale / metrics)
  "leadership" → +500  (led design / architecture using this skill)

Dependencies (already pinned in backend/requirements.txt):
  pymupdf, python-docx, openai.
"""

from __future__ import annotations

import io
import json
import logging
import re
from difflib import get_close_matches
from functools import lru_cache

from app.services.cv_contact import parse_contact
# Re-exported: `cv_parser.normalize_structured` / `.has_content` are the names the
# rest of the codebase already reaches for. The definitions live in a leaf module
# so the repository layer can import the contract without importing the parser.
from app.services.cv_structured_shape import (  # noqa: F401
    BLANK_CONTACT,
    CONTRACT_KEYS,
    coerce_sections,
    has_content,
    normalize_structured,
)
from app.services.cv_explicit_skills import extract_explicit_skills, reconcile_skill_signals
from app.security.personal_data import sanitize_cv_text_for_ai
from app.services.llm_provider import LLMProvider, LLMProviderError, get_llm_provider
from app.services.taxonomy_loader import _name_index, lookup_by_name

# pymupdf (fitz) and python-docx are imported lazily inside the extract_* helpers
# so that unit tests (and environments that don't parse PDFs) can import this
# module without those heavy dependencies being present.

logger = logging.getLogger(__name__)

# ── Config ────────────────────────────────────────────────────────────────────

_SIGNAL_XP: dict[str, int] = {
    "mention":    50,
    "project":    150,
    "impact":     350,
    "leadership": 500,
}
_VALID_SIGNALS: set[str] = set(_SIGNAL_XP.keys())

_MAX_SKILLS = 50
_FUZZY_THRESHOLD = 0.88
_CV_TEXT_CHAR_LIMIT = 15_000  # truncate very long CVs before sending to LLM
# Output budget for the structured re-parse. 15k chars in is roughly 4k tokens,
# and the JSON restating them with keys and escaping is strictly larger — so the
# ceiling has to sit comfortably above the input, not at it.
_STRUCTURED_MAX_OUTPUT_TOKENS = 12_000
_MIN_RAW_TEXT_LEN = 80        # below this we assume scanned / empty CV
_MIN_VOWEL_RATIO = 0.15       # real text ≥ 35%; keyboard-smash is typically < 5%



# ── Input quality guard ───────────────────────────────────────────────────────

def _is_plausible_professional_text(text: str) -> bool:
    """Return False for obvious keyboard-smash — saves LLM quota and gives a
    clear 'add real content' response instead of a false provider-failure 503."""
    letters = re.sub(r"[^a-zA-Z]", "", text)
    if len(letters) < 10:
        return False
    vowels = sum(1 for c in letters.lower() if c in "aeiou")
    return (vowels / len(letters)) >= _MIN_VOWEL_RATIO


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_text_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pymupdf (fitz). Lazy-imported."""
    import fitz  # pymupdf

    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        pages = [page.get_text("text", sort=True) for page in doc]
    return "\n\n".join(pages).strip()


def _extract_text_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx. Walks paragraphs + tables."""
    from docx import Document  # lazy import

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()


# ── LLM extraction ────────────────────────────────────────────────────────────

_SYSTEM_PROMPT = """You are an expert CV analyst. Given a candidate's CV text, return BOTH a comprehensive Lightcast-skill list AND a structured breakdown of the CV's sections.

The candidate's contact header has been removed before you see this text — it is
read locally. Do not ask for it, infer it, or invent one.

Return JSON only — no prose, no markdown fences. Top-level shape:
{
  "skills":     [{"taxonomy_key": "...", "signal_type": "...", "evidence": "..."}],
  "structured": {
    "summary":     "string | null",
    "education":   [{"institution":"...","degree":"...","dates":"...","grade":"...","location":"..."}],
    "experience":  [{"company":"...","role":"...","dates":"...","location":"...","bullets":["...", "..."]}],
    "projects":    [{"name":"...","dates":"...","bullets":["...", "..."]}],
    "skills_line": "string | null",
    "certs":       ["...", "..."]
  }
}

For each skill in "skills":
  - "taxonomy_key": the Lightcast skill name. Prefer canonical forms — e.g. "Python (Programming Language)" not just "Python", "SQL (Programming Language)" not "SQL", "Data Warehousing", "Stakeholder Management".
  - "signal_type": one of: "mention", "project", "impact", "leadership"
      mention    — named in a skills list only, no evidence of use
      project    — used in a real project, role, or product
      impact     — applied with measurable outcome / metrics / scale
      leadership — led design, architecture, or team using this skill
  - "evidence": a short phrase or sentence from the CV that justifies the signal (≤200 chars)

Skill rules:
  - Include hard skills, tools, methodologies, AND human skills when evidenced
  - Skip generic filler ("Innovation", "Collaboration" alone) unless there is concrete evidence tied to a project/outcome
  - If the CV mentions a skill in multiple contexts, use the HIGHEST signal_type
  - Return 20–50 skills. Extract only what is evidenced — do not invent skills
  - Use proper Lightcast capitalisation (e.g. "Apache Spark", "Amazon Web Services (AWS)")

Structured-section rules:
  - "summary": the opening paragraph / objective if present, else null
  - "education": every degree row. Use empty string for missing fields, not omission. dates as printed (e.g. "March 2024", "Jun 2020")
  - "experience": every role. Preserve role order top→bottom of CV. "bullets" = each "•" / "-" / numbered line under that role, verbatim, ≤300 chars each. No bullet-merging. If a role has no bullets, return [].
  - "projects": only true "Projects" / "Personal Projects" sections. Do NOT duplicate role-level work as projects.
  - "skills_line": the single skills paragraph (comma/pipe-separated list) verbatim, or null if no dedicated skills section
  - "certs": each certification as a single string. Empty array if none.
  - Do NOT invent fields, dates, or bullets. Extract only what is present in the CV text."""

_SKILLS_SYSTEM_PROMPT = """Extract every evidenced professional skill from this CV.
Return JSON only as an array of objects with exactly these keys:
[{"taxonomy_key":"canonical Lightcast skill name","signal_type":"mention|project|impact|leadership","evidence":"short verbatim evidence"}]

Rules:
- Include explicit hard skills, tools, methods, and evidenced human skills.
- Prefer canonical Lightcast names such as "Python (Programming Language)".
- mention = listed only; project = used in work/project; impact = measurable outcome; leadership = led team/design/architecture.
- Keep the strongest signal when a skill appears more than once.
- Extract only evidenced skills. Never infer or invent a skill.
- Return at most 50 skills. No prose or markdown fences."""


async def _llm_extract_skills(
    cv_text: str,
    provider: LLMProvider | None = None,
) -> tuple[list[dict] | None, dict[str, object]]:
    """Fast judgment path: skill extraction only, with model provenance."""
    messages = [
        {"role": "system", "content": _SKILLS_SYSTEM_PROMPT},
        {
            "role": "user",
            "content": f"CV text:\n---\n{sanitize_cv_text_for_ai(cv_text)[:_CV_TEXT_CHAR_LIMIT]}\n---\nReturn the JSON array only.",
        },
    ]
    selected = provider or get_llm_provider()
    metadata: dict[str, object] = {}
    try:
        if hasattr(selected, "complete_with_metadata"):
            completion = await selected.complete_with_metadata(
                messages,
                max_tokens=3072,
                temperature=0,
            )
            raw = completion.content
            metadata = {
                "llm_model": completion.model,
                "llm_elapsed_ms": completion.elapsed_ms,
            }
        else:
            raw = await selected.complete(messages, max_tokens=3072, temperature=0)
    except LLMProviderError:
        logger.error("All CV skill extraction providers failed")
        return None, metadata

    skills, _structured = _parse_llm_json(raw)
    if skills is None:
        logger.warning("CV skill extraction returned unparseable JSON")
    return skills, metadata


async def _llm_extract(
    cv_text: str,
    provider: LLMProvider | None = None,
) -> tuple[list[dict] | None, dict | None]:
    """Single LLM call → (skills, cv_structured).

    Returns:
        (list[dict], dict)  — both extracted cleanly
        (list[dict], None)  — skills parsed but structured payload missing/invalid (degraded but usable)
        (None, None)        — every provider failed or returned unparseable output
    """
    truncated = sanitize_cv_text_for_ai(cv_text)[:_CV_TEXT_CHAR_LIMIT]
    messages = [
        {"role": "system", "content": _SYSTEM_PROMPT},
        {"role": "user", "content": (
            f"CV text:\n---\n{truncated}\n---\n\n"
            "Return JSON only with top-level keys 'skills' and 'structured' per the rules."
        )},
    ]
    try:
        raw = await (provider or get_llm_provider()).complete(messages, max_tokens=4096, temperature=0)
    except LLMProviderError:
        logger.error("All CV extraction providers failed")
        return None, None

    skills, structured = _parse_llm_json(raw)
    if skills is None:
        logger.warning("CV extraction: provider responded but returned unparseable JSON")
    elif structured is None:
        logger.info("CV extraction: skills parsed but 'structured' payload missing")
    return skills, attach_contact(structured, cv_text)


def attach_contact(structured: dict | None, raw_text: str) -> dict | None:
    """Fill `structured["contact"]` from the raw CV text, deterministically.

    The single place a parsed CV gets its identity. Every path that produces a
    `cv_structured` payload must run through here, because the model is never
    shown the header and therefore can never supply one.
    """
    if structured is None:
        return None
    structured["contact"] = parse_contact(raw_text)
    return structured


def _parse_llm_json(text: str) -> tuple[list[dict] | None, dict | None]:
    """Pull skills + structured out of an LLM response — tolerates code fences.

    Returns:
        (list[dict], dict | None)  — skills parsed (possibly with structured payload)
        (None, None)               — could not locate or decode valid JSON
    """
    if not text:
        return None, None

    if "```" in text:
        match = re.search(r"```(?:json)?\s*(.+?)```", text, flags=re.DOTALL)
        if match:
            text = match.group(1).strip()

    start_obj = text.find("{")
    start_arr = text.find("[")
    starts = [s for s in (start_obj, start_arr) if s >= 0]
    if not starts:
        return None, None

    try:
        parsed, _ = json.JSONDecoder().raw_decode(text, min(starts))
    except json.JSONDecodeError:
        logger.warning("CV extractor: could not parse LLM response as JSON")
        return None, None

    if isinstance(parsed, list):
        # Legacy bare-list response — skills only, no structure.
        return parsed, None

    if isinstance(parsed, dict):
        skills: list[dict] | None = None
        structured: dict | None = None

        if "skills" in parsed and isinstance(parsed["skills"], list):
            skills = parsed["skills"]
        else:
            # Fallback: accept any list-valued field as skills (legacy clients).
            for v in parsed.values():
                if isinstance(v, list):
                    skills = v
                    break

        if "structured" in parsed and isinstance(parsed["structured"], dict):
            structured = _validate_structured(parsed["structured"])

        return skills, structured

    return None, None


def _validate_structured(raw: dict) -> dict | None:
    """Coerce an LLM structured payload into the contract, with a blank contact.

    `contact` is NOT taken from the model. The header is stripped before the
    prompt is built, so anything the model returns there is a hallucination or —
    as shipped for two weeks — the redaction placeholder itself. Callers holding
    the raw text fill it via `attach_contact`.
    """
    if not isinstance(raw, dict):
        return None
    out = coerce_sections(raw, ingest=True)
    out["contact"] = dict(BLANK_CONTACT)
    return out


# ── Validation / Lightcast mapping ────────────────────────────────────────────

@lru_cache(maxsize=1)
def _lightcast_name_list() -> list[str]:
    """Cached lowercase list of all Lightcast names for fuzzy matching."""
    return list(_name_index().keys())


def _fuzzy_match(name: str) -> str | None:
    """Return the canonical Lightcast name if a close match exists."""
    norm = name.lower().strip()
    if not norm:
        return None
    candidates = get_close_matches(norm, _lightcast_name_list(), n=1, cutoff=_FUZZY_THRESHOLD)
    if not candidates:
        return None
    return _name_index()[candidates[0]].name


def _validate_and_normalize(raw_skills: list[dict]) -> list[dict]:
    """
    Keep only skills whose taxonomy_key resolves to a real Lightcast name.
    - Coerce signal_type to a known value (defaults to "mention")
    - Compute xp_awarded from signal_type
    - Dedupe by canonical name, keeping the highest-signal occurrence
    - Cap at _MAX_SKILLS, sorted by xp descending
    """
    best_by_key: dict[str, dict] = {}

    for raw in raw_skills[: _MAX_SKILLS * 2]:
        if not isinstance(raw, dict):
            continue

        name = str(raw.get("taxonomy_key") or raw.get("name") or "").strip()
        signal = str(raw.get("signal_type") or "mention").strip().lower()
        evidence = str(raw.get("evidence") or "").strip()

        if not name:
            continue
        if signal not in _VALID_SIGNALS:
            signal = "mention"

        # Resolve to canonical Lightcast name
        lc = lookup_by_name(name)
        canonical = lc.name if lc else _fuzzy_match(name)
        if canonical is None:
            logger.info("Dropped one non-Lightcast skill from CV extraction")
            continue

        new_xp = _SIGNAL_XP[signal]
        prev = best_by_key.get(canonical)
        if prev is None or prev["xp_awarded"] < new_xp:
            best_by_key[canonical] = {
                "taxonomy_key": canonical,
                "xp_awarded": new_xp,
                "signal_type": signal,
                "evidence": evidence[:300],
            }

    return sorted(best_by_key.values(), key=lambda s: -s["xp_awarded"])[:_MAX_SKILLS]


# ── Main entry points ─────────────────────────────────────────────────────────

def extract_raw_text(file_bytes: bytes, file_type: str) -> str:
    """Extract plain text from a CV file without running LLM analysis."""
    if file_type == "pdf":
        return _extract_text_pdf(file_bytes)
    if file_type == "docx":
        return _extract_text_docx(file_bytes)
    raise ValueError(f"Unsupported file_type: {file_type!r}")


async def parse_cv_text(raw_text: str, provider: LLMProvider | None = None) -> dict:
    """Parse free-form self-description text and return skill signals + structured sections.

    Result keys:
        skills_detected  — validated skill list (may be empty)
        cv_structured    — {summary, education[], experience[], projects[], skills_line, certs[]} or None
        raw_text         — original input
        provider_failed  — True when every LLM provider errored or returned
                           unparseable output; False when at least one responded
                           with valid JSON (even if skills list is empty)
    """
    if not raw_text or len(raw_text.strip()) < _MIN_RAW_TEXT_LEN:
        return {"skills_detected": [], "cv_structured": None, "raw_text": raw_text, "provider_failed": False}

    if not _is_plausible_professional_text(raw_text):
        logger.info("CV text rejected as non-linguistic input (%d chars)", len(raw_text))
        return {"skills_detected": [], "cv_structured": None, "raw_text": raw_text, "provider_failed": False}

    # Literal recall first, exactly as the persisted-CV path does.
    #
    # This path used to trust the model alone, so the SAME CV could yield fewer
    # skills — or none at all — before signup than after it. A CV naming taxonomy
    # skills outright would 422 with "we couldn't pull skills from this CV" on the
    # free preview and then succeed once uploaded, because only `parse_cv_skills`
    # reconciled the deterministic signal. That asymmetry sat on the widest step of
    # the funnel, where most people who never come back are lost.
    deterministic = extract_explicit_skills(raw_text)
    raw_skills, structured = await _llm_extract(raw_text, provider)
    enriched = _validate_and_normalize(raw_skills or [])
    skills = reconcile_skill_signals(deterministic, enriched)
    # A model outage with real deterministic skills in hand is a degraded read,
    # not a failure — callers fork on this to decide 503 vs. serve what we have.
    provider_failed = raw_skills is None and not skills
    logger.info(
        "Self-description parsed: %d chars → %d deterministic + %d raw → %d validated Lightcast skills (structured=%s, provider_failed=%s)",
        len(raw_text), len(deterministic), len(raw_skills or []), len(skills),
        structured is not None, provider_failed,
    )
    return {
        "skills_detected": skills,
        "cv_structured":   structured,
        "raw_text":        raw_text,
        "provider_failed": provider_failed,
    }


async def parse_cv_skills(raw_text: str, provider: LLMProvider | None = None) -> dict:
    """Trusted persisted-CV path: literal recall plus model enrichment.

    Explicit taxonomy skills are deterministic, so an incomplete model response
    can enrich but can never erase what the candidate actually wrote.
    """
    deterministic = extract_explicit_skills(raw_text)
    if not raw_text or len(raw_text.strip()) < _MIN_RAW_TEXT_LEN:
        return {
            "skills_detected": deterministic,
            "cv_structured": None,
            "raw_text": raw_text,
            "provider_failed": not bool(deterministic),
            "provenance": {
                "deterministic_skill_count": len(deterministic),
                "llm_skill_count": 0,
                "validated_skill_count": len(deterministic),
            },
        }

    raw_skills, model_metadata = await _llm_extract_skills(raw_text, provider)
    enriched = _validate_and_normalize(raw_skills or [])
    skills = reconcile_skill_signals(deterministic, enriched)
    llm_failed = raw_skills is None
    return {
        "skills_detected": skills,
        "cv_structured": None,
        "raw_text": raw_text,
        "provider_failed": llm_failed and not skills,
        "llm_enrichment_failed": llm_failed,
        "provenance": {
            "deterministic_skill_count": len(deterministic),
            "llm_skill_count": len(raw_skills or []),
            "validated_llm_skill_count": len(enriched),
            "validated_skill_count": len(skills),
            **model_metadata,
        },
    }


async def parse_cv(file_bytes: bytes, file_type: str, provider: LLMProvider | None = None) -> dict:
    """
    Parse a CV and return detected skill signals mapped to Lightcast.

    Args:
        file_bytes: Raw bytes of the uploaded PDF or DOCX.
        file_type:  "pdf" or "docx".

    Returns:
        {"skills_detected": [
             {"taxonomy_key", "xp_awarded", "signal_type", "evidence"}
         ],
         "raw_text": str}

        Empty `skills_detected` signals to the caller that the file is likely
        scanned/empty or extraction failed — router should translate to HTTP 422.

    Raises:
        ValueError — unsupported file_type.
    """
    if file_type == "pdf":
        raw_text = _extract_text_pdf(file_bytes)
    elif file_type == "docx":
        raw_text = _extract_text_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file_type: {file_type!r}")

    if len(raw_text) < _MIN_RAW_TEXT_LEN:
        logger.warning("CV extracted text is too short (%d chars) — likely scanned", len(raw_text))
        return {"skills_detected": [], "cv_structured": None, "raw_text": raw_text, "provider_failed": False}

    raw_skills, structured = await _llm_extract(raw_text, provider)
    provider_failed = raw_skills is None
    skills = _validate_and_normalize(raw_skills or [])
    logger.info(
        "CV parsed: %d chars → %d raw → %d validated Lightcast skills (structured=%s, provider_failed=%s)",
        len(raw_text), len(raw_skills or []), len(skills), structured is not None, provider_failed,
    )

    return {
        "skills_detected": skills,
        "cv_structured":   structured,
        "raw_text":        raw_text,
        "provider_failed": provider_failed,
    }


# ── Structured re-parse (lazy backfill) ───────────────────────────────────────

_STRUCTURED_ONLY_PROMPT = """You are an expert CV analyst. Given a candidate's CV text, return ONLY the structured section breakdown — no skills extraction.

Return JSON only — no prose, no markdown fences. Shape:
{
  "summary":     "string | null",
  "education":   [{"institution":"...","degree":"...","dates":"...","grade":"...","location":"..."}],
  "experience":  [{"company":"...","role":"...","dates":"...","location":"...","bullets":["...", "..."]}],
  "projects":    [{"name":"...","dates":"...","bullets":["...", "..."]}],
  "skills_line": "string | null",
  "certs":       ["...", "..."]
}

Rules:
  - "summary": opening paragraph / objective if present, else null
  - "education": every degree row. Use empty string for missing fields.
  - "experience": every role. Preserve order. "bullets" = each "•" / "-" / numbered line under that role, verbatim, ≤300 chars each. No bullet-merging.
  - "projects": only true "Projects" sections — do NOT duplicate role-level work.
  - "skills_line": dedicated skills paragraph verbatim, or null
  - "certs": each certification as a single string
  - Do NOT invent fields, dates, or bullets."""


async def reparse_structured_only(raw_text: str) -> dict | None:
    """Re-parse existing cv_raw_text to fill cv_structured (lazy backfill).

    Returns the validated structured payload or None on provider failure.
    """
    if not raw_text or len(raw_text.strip()) < _MIN_RAW_TEXT_LEN:
        return None

    truncated = sanitize_cv_text_for_ai(raw_text)[:_CV_TEXT_CHAR_LIMIT]
    messages = [
        {"role": "system", "content": _STRUCTURED_ONLY_PROMPT},
        {"role": "user", "content": (
            f"CV text:\n---\n{truncated}\n---\n\n"
            "Return JSON only with the structured sections per the rules."
        )},
    ]
    try:
        # 4096 was too tight for the job it is given. The prompt asks for every
        # bullet of every role VERBATIM from up to _CV_TEXT_CHAR_LIMIT (15k) chars
        # of input, so a dense CV's JSON does not fit in 4k output tokens and
        # arrives truncated — which reads downstream as "unparseable JSON", i.e.
        # as a provider fault rather than a budget we set. Output must be able to
        # exceed the input it is restructuring.
        raw = await get_llm_provider().complete(
            messages, max_tokens=_STRUCTURED_MAX_OUTPUT_TOKENS, temperature=0
        )
    except LLMProviderError:
        logger.error("Structured re-parse: all providers failed")
        return None

    if not raw:
        return None

    if "```" in raw:
        match = re.search(r"```(?:json)?\s*(.+?)```", raw, flags=re.DOTALL)
        if match:
            raw = match.group(1).strip()

    start = raw.find("{")
    if start < 0:
        logger.warning(
            "Structured re-parse: no JSON object in %d-char response; head=%r",
            len(raw), raw[:200],
        )
        return None
    try:
        parsed, _ = json.JSONDecoder().raw_decode(raw, start)
    except json.JSONDecodeError:
        # Log enough to tell the two causes apart without a second incident.
        # A response that simply STOPS mid-structure is the output budget running
        # out; one that ends with prose or a stray fence is the model ignoring the
        # format. "unparseable JSON" alone forced that call to be guesswork.
        logger.warning(
            "Structured re-parse: unparseable JSON (%d chars, looks_truncated=%s) tail=%r",
            len(raw),
            not raw.rstrip().endswith("}"),
            raw[-200:],
        )
        return None

    if not isinstance(parsed, dict):
        return None
    return attach_contact(_validate_structured(parsed), raw_text)
