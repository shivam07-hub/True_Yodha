"""Structured CV → DOCX bytes using python-docx.

Builds an ATS-clean Word document from the SAME visible sections the frontend
renders on the `.cvb-pdf-page` sheet (selectVisibleCV applied client-side), so
the .docx never diverges from the WYSIWYG PDF. Single column, real headings,
real bullet lists, selectable text — exactly what an applicant tracking system
parses cleanly.

python-docx is already a backend dependency (used for DOCX text extraction);
this adds the generation direction. No headless browser, no new infra.
"""
from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor

from app.services.cv_section_order import normalize_section_order

_DARK = RGBColor(0x0A, 0x0A, 0x0A)
_BODY = RGBColor(0x1A, 0x1A, 0x1A)
_MUTED = RGBColor(0x55, 0x55, 0x55)

# Per-template type scale. All ATS-safe; compact just condenses to fit one page.
_SCALE: dict[str, dict[str, float]] = {
    "classic": {"name": 20, "section": 10.5, "role": 11, "body": 10.5, "space_after": 4},
    "modern":  {"name": 18, "section": 10.5, "role": 11, "body": 10.5, "space_after": 4},
    "compact": {"name": 17, "section": 9.5, "role": 10, "body": 9.5, "space_after": 2},
}


def _scale(template: str) -> dict[str, float]:
    return _SCALE.get(template, _SCALE["classic"])


def generate_cv_docx(visible: dict[str, Any], contact: dict[str, Any], template: str = "classic") -> bytes:
    """Render the visible CV model + contact into ATS-clean .docx bytes."""
    s = _scale(template)
    doc = Document()

    # Base document font.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(s["body"])
    normal.font.color.rgb = _BODY

    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    # ── Header: name, title, contact line ─────────────────────────────
    name = (contact.get("name") or "").strip() or "Your Name"
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(s["name"])
    name_run.font.color.rgb = _DARK

    title = (contact.get("title") or "").strip()
    if title:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(6)
        tr = tp.add_run(title)
        tr.font.size = Pt(s["body"] + 0.5)
        tr.font.color.rgb = _MUTED

    contact_bits = [
        (contact.get("location") or "").strip(),
        (contact.get("email") or "").strip(),
        (contact.get("phone") or "").strip(),
        (contact.get("linkedin") or "").strip(),
    ]
    contact_line = "  ·  ".join(b for b in contact_bits if b)
    if contact_line:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(8)
        cr = cp.add_run(contact_line)
        cr.font.size = Pt(s["body"] - 1.5)
        cr.font.color.rgb = _MUTED

    def heading(text: str) -> None:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(8)
        hp.paragraph_format.space_after = Pt(3)
        hr = hp.add_run(text.upper())
        hr.bold = True
        hr.font.size = Pt(s["section"])
        hr.font.color.rgb = _DARK

    def role_line(left_bold: str, left_rest: str, right: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        if right:
            p.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
        rb = p.add_run(left_bold)
        rb.bold = True
        rb.font.size = Pt(s["role"])
        rb.font.color.rgb = _DARK
        if left_rest:
            rr = p.add_run(left_rest)
            rr.font.size = Pt(s["body"])
            rr.font.color.rgb = _MUTED
        if right:
            dr = p.add_run("\t" + right)
            dr.font.size = Pt(s["body"] - 1)
            dr.font.color.rgb = _MUTED

    def bullets(items: list[str]) -> None:
        for b in items:
            if not (b or "").strip():
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(s["space_after"] - 2 if s["space_after"] > 2 else 0)
            br = bp.add_run(b.strip())
            br.font.size = Pt(s["body"])
            br.font.color.rgb = _BODY

    def emit_summary() -> None:
        summary = (visible.get("summary") or "").strip()
        if not summary:
            return
        heading("Summary")
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(s["space_after"])
        sr = sp.add_run(summary)
        sr.font.size = Pt(s["body"])
        sr.font.color.rgb = _BODY

    def emit_experience() -> None:
        experience = [e for e in (visible.get("experience") or []) if (e.get("bullets") or [])]
        if not experience:
            return
        heading("Experience")
        for e in experience:
            company = (e.get("company") or "").strip()
            role_line(e.get("role") or "", f"  ·  {company}" if company else "", (e.get("dates") or "").strip())
            bullets(e.get("bullets") or [])

    def emit_projects() -> None:
        projects = [p for p in (visible.get("projects") or []) if (p.get("bullets") or [])]
        if not projects:
            return
        heading("Projects")
        for pr in projects:
            role_line(pr.get("name") or "", "", (pr.get("dates") or "").strip())
            bullets(pr.get("bullets") or [])

    def emit_education() -> None:
        education = visible.get("education") or []
        if not education:
            return
        heading("Education")
        for ed in education:
            degree = (ed.get("degree") or "").strip()
            grade = (ed.get("grade") or "").strip()
            sub = "".join(f"  ·  {part}" for part in (degree, grade) if part)
            role_line(ed.get("institution") or "", sub, (ed.get("dates") or "").strip())

    def emit_skills() -> None:
        skills_line = (visible.get("skills_line") or "").strip()
        if not skills_line:
            return
        heading("Skills")
        kp = doc.add_paragraph()
        kp.paragraph_format.space_after = Pt(s["space_after"])
        kr = kp.add_run(skills_line)
        kr.font.size = Pt(s["body"])
        kr.font.color.rgb = _BODY

    def emit_certs() -> None:
        certs = [c for c in (visible.get("certs") or []) if (c or "").strip()]
        if not certs:
            return
        heading("Certifications")
        bullets(certs)

    emitters = {
        "summary": emit_summary,
        "experience": emit_experience,
        "projects": emit_projects,
        "skills_line": emit_skills,
        "education": emit_education,
        "certs": emit_certs,
    }
    for key in normalize_section_order(visible.get("order")):
        emitters[key]()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
